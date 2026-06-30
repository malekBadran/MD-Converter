"""Render a parsed Markdown token stream into a python-docx Document.

RTL/Arabic handling: python-docx's high-level API has no notion of
bidirectional text, so this module talks to the underlying OOXML directly
(`w:bidi` on paragraphs, `w:rtl`/`w:cs` + `w:rFonts/@w:cs` on runs). Text is
split into script-homogeneous runs (see md2doc.lang.split_bidi_runs) so a
single paragraph or table cell can correctly mix Arabic and Latin/numeric
text, each part getting the right direction and font.
"""

from __future__ import annotations

import io
import urllib.error
import urllib.request
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it.token import Token
from PIL import Image

from .lang import is_rtl_text, split_bidi_runs
from .mermaid import MermaidRenderer
from .parser import find_block_end

RTL_THRESHOLD = 0.3
CODE_FONT = "Consolas"
MAX_IMAGE_WIDTH_IN = 6.2
_USER_AGENT = "md2doc/1.0 (+https://github.com)"


@dataclass
class DocxOptions:
    arabic_font: str = "Arial"
    latin_font: str = "Calibri"
    allow_network: bool = True
    base_dir: Path = field(default_factory=Path.cwd)
    warn: callable = lambda msg: None


# ---------------------------------------------------------------------------
# Low-level OOXML helpers (bidi / complex-script support python-docx lacks)
# ---------------------------------------------------------------------------

def _set_paragraph_bidi(paragraph, rtl: bool) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:bidi"))
    if rtl and existing is None:
        pPr.append(OxmlElement("w:bidi"))
    elif not rtl and existing is not None:
        pPr.remove(existing)


def _set_table_bidi(table, rtl: bool) -> None:
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn("w:bidiVisual"))
    if rtl and existing is None:
        tblPr.append(OxmlElement("w:bidiVisual"))
    elif not rtl and existing is not None:
        tblPr.remove(existing)


def _style_run(run, rtl: bool, arabic_font: str, latin_font: str, code: bool = False) -> None:
    rPr = run._r.get_or_add_rPr()
    if rtl:
        if rPr.find(qn("w:rtl")) is None:
            rPr.append(OxmlElement("w:rtl"))
        if rPr.find(qn("w:cs")) is None:
            rPr.append(OxmlElement("w:cs"))
    rFonts = rPr.get_or_add_rFonts()
    font_name = CODE_FONT if code else (arabic_font if rtl else latin_font)
    rFonts.set(qn("w:cs"), font_name)
    if not code:
        run.font.name = latin_font if not rtl else run.font.name or latin_font
    else:
        run.font.name = CODE_FONT
        rFonts.set(qn("w:ascii"), CODE_FONT)
        rFonts.set(qn("w:hAnsi"), CODE_FONT)


def _shade_paragraph(paragraph, color_hex: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def _add_bottom_border(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B0B3B8")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Inline content (bold/italic/code/links/images/line breaks) -> runs
# ---------------------------------------------------------------------------

class InlineStyle:
    __slots__ = ("bold", "italic", "strike", "code", "link")

    def __init__(self):
        self.bold = False
        self.italic = False
        self.strike = False
        self.code = False
        self.link = None


def _add_styled_text(paragraph, text: str, style: InlineStyle, opts: DocxOptions) -> None:
    if not text:
        return
    for chunk, is_rtl in split_bidi_runs(text):
        run = paragraph.add_run(chunk)
        run.bold = style.bold
        run.italic = style.italic
        run.font.strike = style.strike
        if style.code:
            run.font.name = CODE_FONT
            run.font.size = Pt(10)
        if style.link:
            run.font.underline = True
            run.font.color.rgb = RGBColor(0x09, 0x69, 0xDA)
        _style_run(run, is_rtl, opts.arabic_font, opts.latin_font, code=style.code)


def _resolve_image_bytes(src: str, opts: DocxOptions) -> bytes | None:
    try:
        if src.startswith("http://") or src.startswith("https://"):
            if not opts.allow_network:
                opts.warn(f"Skipping remote image (network disabled): {src}")
                return None
            req = urllib.request.Request(src, headers={"User-Agent": _USER_AGENT})
            with urllib.request.urlopen(req, timeout=20) as resp:
                return resp.read()
        path = (opts.base_dir / src).resolve()
        if not path.exists():
            opts.warn(f"Image not found, skipping: {src}")
            return None
        return path.read_bytes()
    except (urllib.error.URLError, OSError) as exc:
        opts.warn(f"Failed to load image '{src}': {exc}")
        return None


def _add_image_run(paragraph, data: bytes, opts: DocxOptions) -> None:
    try:
        with Image.open(io.BytesIO(data)) as im:
            width_px, height_px = im.size
        width_in = min(MAX_IMAGE_WIDTH_IN, width_px / 96)
        height_in = height_px / width_px * width_in if width_px else None
    except Exception:
        width_in, height_in = MAX_IMAGE_WIDTH_IN, None
    run = paragraph.add_run()
    if height_in:
        run.add_picture(io.BytesIO(data), width=Inches(width_in), height=Inches(height_in))
    else:
        run.add_picture(io.BytesIO(data), width=Inches(width_in))


def render_inline(paragraph, children: list[Token], opts: DocxOptions) -> None:
    style = InlineStyle()
    stack: list[str] = []

    for tok in children:
        t = tok.type
        if t == "text":
            _add_styled_text(paragraph, tok.content, style, opts)
        elif t == "softbreak":
            _add_styled_text(paragraph, " ", style, opts)
        elif t == "hardbreak":
            paragraph.add_run().add_break()
        elif t == "code_inline":
            style.code = True
            _add_styled_text(paragraph, tok.content, style, opts)
            style.code = False
        elif t in ("strong_open",):
            style.bold = True
            stack.append("bold")
        elif t == "strong_close":
            style.bold = False
        elif t == "em_open":
            style.italic = True
            stack.append("italic")
        elif t == "em_close":
            style.italic = False
        elif t == "s_open":
            style.strike = True
            stack.append("strike")
        elif t == "s_close":
            style.strike = False
        elif t == "link_open":
            style.link = tok.attrGet("href")
        elif t == "link_close":
            style.link = None
        elif t == "image":
            alt = tok.content or ""
            src = tok.attrGet("src") or ""
            data = _resolve_image_bytes(src, opts)
            if data:
                _add_image_run(paragraph, data, opts)
            else:
                _add_styled_text(paragraph, f"[image: {alt or src}]", style, opts)
        # unknown inline tokens are silently skipped (defensive: unsupported syntax)


def _paragraph_text(children: list[Token]) -> str:
    return "".join(t.content for t in children if t.type in ("text", "code_inline"))


# ---------------------------------------------------------------------------
# Block-level rendering
# ---------------------------------------------------------------------------

HEADING_STYLE = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4", 5: "Heading 5", 6: "Heading 6"}
BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]
NUMBER_STYLES = ["List Number", "List Number 2", "List Number 3"]


class DocxBuilder:
    def __init__(self, opts: DocxOptions, mermaid_renderer: MermaidRenderer | None):
        self.doc = Document()
        self.opts = opts
        self.mermaid = mermaid_renderer
        self._setup_styles()

    def _setup_styles(self) -> None:
        normal = self.doc.styles["Normal"]
        normal.font.name = self.opts.latin_font
        normal.font.size = Pt(11)

    def build(self, tokens: list[Token]) -> Document:
        idx = 0
        while idx < len(tokens):
            idx = self._dispatch(tokens, idx)
        return self.doc

    def _dispatch(self, tokens: list[Token], idx: int) -> int:
        tok = tokens[idx]
        if tok.type == "heading_open":
            return self._render_heading(tokens, idx)
        if tok.type == "paragraph_open":
            return self._render_paragraph(tokens, idx)
        if tok.type in ("bullet_list_open", "ordered_list_open"):
            return self._render_list(tokens, idx, level=0)
        if tok.type == "blockquote_open":
            return self._render_blockquote(tokens, idx)
        if tok.type == "table_open":
            return self._render_table(tokens, idx)
        if tok.type == "fence":
            self._render_fence(tok)
            return idx + 1
        if tok.type == "code_block":
            self._render_code_block(tok.content)
            return idx + 1
        if tok.type == "hr":
            self._render_hr()
            return idx + 1
        if tok.type in ("html_block",):
            return idx + 1
        return idx + 1

    # -- headings / paragraphs --------------------------------------------

    def _render_heading(self, tokens: list[Token], idx: int) -> int:
        open_tok = tokens[idx]
        level = int(open_tok.tag[1])
        inline = tokens[idx + 1]
        p = self.doc.add_paragraph(style=HEADING_STYLE.get(level, "Heading 6"))
        self._apply_direction(p, inline.content)
        render_inline(p, inline.children or [], self.opts)
        return find_block_end(tokens, idx) + 1

    def _render_paragraph(self, tokens: list[Token], idx: int, base_style: str | None = None) -> int:
        inline = tokens[idx + 1]
        p = self.doc.add_paragraph(style=base_style)
        self._apply_direction(p, inline.content)
        render_inline(p, inline.children or [], self.opts)
        return find_block_end(tokens, idx) + 1

    def _apply_direction(self, paragraph, text: str) -> None:
        rtl = is_rtl_text(text, RTL_THRESHOLD)
        _set_paragraph_bidi(paragraph, rtl)
        if rtl:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # -- lists --------------------------------------------------------------

    def _render_list(self, tokens: list[Token], idx: int, level: int) -> int:
        ordered = tokens[idx].type == "ordered_list_open"
        end = find_block_end(tokens, idx)
        i = idx + 1
        while i < end:
            tok = tokens[i]
            if tok.type == "list_item_open":
                item_end = find_block_end(tokens, i)
                i = self._render_list_item(tokens, i, item_end, level, ordered)
            else:
                i += 1
        return end + 1

    def _render_list_item(self, tokens: list[Token], idx: int, end: int, level: int, ordered: bool) -> int:
        styles = NUMBER_STYLES if ordered else BULLET_STYLES
        style_name = styles[min(level, len(styles) - 1)]
        i = idx + 1
        first_para = True
        while i < end:
            tok = tokens[i]
            if tok.type == "paragraph_open":
                inline = tokens[i + 1]
                p = self.doc.add_paragraph(style=style_name if first_para else None)
                if not first_para:
                    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
                self._apply_direction(p, inline.content)
                render_inline(p, inline.children or [], self.opts)
                first_para = False
                i = find_block_end(tokens, i) + 1
            elif tok.type in ("bullet_list_open", "ordered_list_open"):
                i = self._render_list(tokens, i, level + 1)
            elif tok.type == "fence":
                self._render_fence(tok)
                i += 1
            else:
                i += 1
        return end + 1

    # -- blockquote -----------------------------------------------------

    def _render_blockquote(self, tokens: list[Token], idx: int) -> int:
        end = find_block_end(tokens, idx)
        i = idx + 1
        while i < end:
            tok = tokens[i]
            if tok.type == "paragraph_open":
                inline = tokens[i + 1]
                p = self.doc.add_paragraph(style="Quote")
                p.paragraph_format.left_indent = Inches(0.3)
                self._apply_direction(p, inline.content)
                render_inline(p, inline.children or [], self.opts)
                i = find_block_end(tokens, i) + 1
            elif tok.type == "blockquote_open":
                i = self._render_blockquote(tokens, i)
            else:
                i += 1
        return end + 1

    # -- code / mermaid fences ------------------------------------------

    def _render_code_block(self, code: str) -> None:
        p = self.doc.add_paragraph()
        _shade_paragraph(p, "F6F8FA")
        lines = code.rstrip("\n").split("\n")
        for n, line in enumerate(lines):
            run = p.add_run(line if line else " ")
            run.font.name = CODE_FONT
            run.font.size = Pt(10)
            if n != len(lines) - 1:
                run.add_break()

    def _render_fence(self, tok: Token) -> None:
        lang = tok.info.strip().split()[0] if tok.info.strip() else ""
        if lang.lower() == "mermaid" and self.mermaid is not None:
            diagram = self.mermaid.render(tok.content, fmt="png")
            if diagram is not None:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_image_run(p, diagram.data, self.opts)
                return
            self.opts.warn("Embedding raw Mermaid source (render unavailable).")
        self._render_code_block(tok.content)

    # -- horizontal rule --------------------------------------------------

    def _render_hr(self) -> None:
        p = self.doc.add_paragraph()
        _add_bottom_border(p)

    # -- tables -------------------------------------------------------------

    def _render_table(self, tokens: list[Token], idx: int) -> int:
        end = find_block_end(tokens, idx)
        rows: list[tuple[str, list[Token]]] = []  # ('th'|'td', [cell tokens...]) per row -> list of cells
        aggregate_text_parts: list[str] = []

        i = idx + 1
        current_row: list[tuple[str, Token, str]] | None = None  # (tag, inline_token, align)
        table_rows: list[list[tuple[str, Token, str]]] = []
        while i < end:
            tok = tokens[i]
            if tok.type == "tr_open":
                current_row = []
            elif tok.type in ("th_open", "td_open"):
                align = "left"
                style_attr = tok.attrGet("style") or ""
                if "right" in style_attr:
                    align = "right"
                elif "center" in style_attr:
                    align = "center"
                inline = tokens[i + 1]
                current_row.append((tok.tag, inline, align))
                aggregate_text_parts.append(inline.content)
            elif tok.type == "tr_close":
                table_rows.append(current_row)
                current_row = None
            i += 1

        if not table_rows:
            return end + 1

        n_rows = len(table_rows)
        n_cols = max(len(r) for r in table_rows)
        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        table_rtl = is_rtl_text(" ".join(aggregate_text_parts), RTL_THRESHOLD)
        _set_table_bidi(table, table_rtl)

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }

        for r, row in enumerate(table_rows):
            for c, (tag, inline, align) in enumerate(row):
                cell = table.cell(r, c)
                p = cell.paragraphs[0]
                cell_rtl = is_rtl_text(inline.content, RTL_THRESHOLD)
                _set_paragraph_bidi(p, cell_rtl)
                p.alignment = align_map["right" if cell_rtl and align == "left" else align]
                render_inline(p, inline.children or [], self.opts)
                if tag == "th":
                    for run in p.runs:
                        run.bold = True
        return end + 1


def convert_to_docx(
    tokens: list[Token],
    output_path: Path,
    base_dir: Path,
    mermaid_renderer: MermaidRenderer | None,
    arabic_font: str = "Arial",
    allow_network: bool = True,
    warn=None,
) -> None:
    opts = DocxOptions(
        arabic_font=arabic_font,
        base_dir=base_dir,
        allow_network=allow_network,
        warn=warn or (lambda msg: None),
    )
    builder = DocxBuilder(opts, mermaid_renderer)
    doc = builder.build(tokens)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
