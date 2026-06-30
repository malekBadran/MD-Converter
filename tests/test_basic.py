"""Lightweight smoke tests (plain asserts, no test framework dependency).

Run with: python -m tests.test_basic
"""

from __future__ import annotations

import io
import sys
import tempfile
import zipfile
from pathlib import Path

import docx

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from md2doc.docx_writer import convert_to_docx
from md2doc.html_writer import extract_title, render_html
from md2doc.lang import is_rtl_text, rtl_ratio, split_bidi_runs
from md2doc.mermaid import MermaidRenderer
from md2doc.parser import parse

SAMPLE = """# Title عنوان

English paragraph.

نص عربي بالكامل هنا للاختبار.

Mixed سلام text 123 example.

| A | B |
|---|---|
| 1 | عربي |

```mermaid
graph TD; A-->B;
```
"""


def test_lang_detection():
    assert rtl_ratio("hello") == 0.0
    assert rtl_ratio("مرحبا") == 1.0
    assert is_rtl_text("نص عربي بالكامل")
    assert not is_rtl_text("just english")
    runs = split_bidi_runs("Hello سلام 123")
    assert runs[0] == ("Hello ", False)
    assert runs[1][1] is True
    print("OK: lang detection")


def test_html_output():
    tokens = parse(SAMPLE)
    title = extract_title(tokens, fallback="doc")
    assert title == "Title عنوان"
    html = render_html(tokens, title=title, mermaid_renderer=None, static_mermaid=False)
    assert "<!DOCTYPE html>" in html
    assert 'dir="rtl"' in html  # the fully-Arabic paragraph
    assert '<pre class="mermaid">' in html
    assert "<table" in html
    print("OK: html writer")


def test_docx_output():
    tokens = parse(SAMPLE)
    renderer = MermaidRenderer(allow_network=True, on_warning=lambda m: None)
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "out.docx"
        convert_to_docx(
            tokens, out, base_dir=Path(tmp), mermaid_renderer=renderer,
            arabic_font="Arial", allow_network=True, warn=lambda m: None,
        )
        assert out.exists()
        doc = docx.Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert any("عنوان" in t for t in texts)
        assert len(doc.tables) == 1
        cell_texts = [c.text for row in doc.tables[0].rows for c in row.cells]
        assert "عربي" in cell_texts
        # at least one run in the document should carry the w:rtl flag
        found_rtl_run = any(
            "w:rtl" in run._r.xml
            for p in doc.paragraphs
            for run in p.runs
        )
        assert found_rtl_run
    print("OK: docx writer")


def test_docx_is_valid_zip():
    tokens = parse("# Hi\n\nJust english text.\n")
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "out.docx"
        convert_to_docx(
            tokens, out, base_dir=Path(tmp), mermaid_renderer=None,
            allow_network=False, warn=lambda m: None,
        )
        with zipfile.ZipFile(out) as z:
            assert "word/document.xml" in z.namelist()
    print("OK: docx is a valid zip/OOXML package")


if __name__ == "__main__":
    test_lang_detection()
    test_html_output()
    test_docx_output()
    test_docx_is_valid_zip()
    print("All smoke tests passed.")
